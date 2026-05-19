from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('ShopModern E-Commerce Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document provides a high-level overview of the Django E-Commerce architecture, explaining the project flow, folder structures, and how files interact with one another.')
    
    # Section 1: Project Flow
    doc.add_heading('1. High-Level Project Flow', level=1)
    doc.add_paragraph(
        "1. User Requests a Page: A user opens a browser and navigates to the website (e.g., the homepage).\n"
        "2. URL Routing (urls.py): Django receives the request and checks `ecommerce/urls.py` and `store/urls.py` to find which Python function (View) should handle the request.\n"
        "3. Views (views/): The matched view function is called. The view acts as the 'middleman'. It contains the core business logic.\n"
        "4. Models (models/): If the view needs data (like fetching products or saving an order), it interacts with the Models. Models communicate directly with the database using Python objects.\n"
        "5. Templates (templates/): Once the view has the necessary data, it passes that data into an HTML Template. The template dynamically renders the data (like looping through products) to build the final HTML page.\n"
        "6. Response: The fully built HTML page is sent back to the user's browser."
    )
    
    # Section 2: Folder Structure
    doc.add_heading('2. Folder & File Breakdown', level=1)
    
    # Root
    doc.add_heading('Root Directory (`ecommerce/`)', level=2)
    p = doc.add_paragraph()
    p.add_run('manage.py: ').bold = True
    p.add_run("The command-line utility for Django. It's used to run the server (`runserver`), create database tables (`migrate`), and create superusers.")
    
    # Project Settings
    doc.add_heading('Project Settings (`ecommerce/ecommerce/`)', level=2)
    p = doc.add_paragraph()
    p.add_run('settings.py: ').bold = True
    p.add_run("The central configuration file. It controls database connections, installed apps (like our 'store' app), template locations, static files, and media file paths.")
    
    p = doc.add_paragraph()
    p.add_run('urls.py: ').bold = True
    p.add_run("The global routing file. It directs traffic to the admin panel and forwards all other traffic to the `store` app's URLs.")
    
    # Store App
    doc.add_heading("Store Application (`ecommerce/store/`)", level=2)
    doc.add_paragraph("This is where the actual e-commerce logic lives. It is broken down into modular components.")
    
    # Models
    doc.add_heading('Database Models (`store/models/`)', level=3)
    doc.add_paragraph("These files define the database tables. Each class represents a table, and each attribute is a column.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('category.py: ').bold = True
    p.add_run("Defines the Category model (e.g., Electronics, Clothing).")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('customer.py: ').bold = True
    p.add_run("Defines the Customer profile, linked 1-to-1 with standard Django Users to store extra info like phone numbers.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('products.py: ').bold = True
    p.add_run("Defines the Product model, containing price, description, image, and a ForeignKey linking it to a Category.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('orders.py: ').bold = True
    p.add_run("Defines the Order model, linking a Customer to a Product, storing the quantity, price, and shipping details.")
    
    # Views
    doc.add_heading('Business Logic (`store/views/`)', level=3)
    doc.add_paragraph("These files handle the logic for each specific webpage.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('home.py: ').bold = True
    p.add_run("Fetches all products (or filters by category) and handles 'Add to Cart' functionality using Django sessions.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('auth.py: ').bold = True
    p.add_run("Handles the logic for User Registration, Login, and Logout.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('cart.py: ').bold = True
    p.add_run("Reads the user's session data to display the items they've added to their shopping cart.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('checkout.py: ').bold = True
    p.add_run("Processes the cart data, creates official Order records in the database, and clears the cart.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('orders.py: ').bold = True
    p.add_run("Fetches a logged-in user's past orders to display on the Order History page.")
    
    # Templates
    doc.add_heading('HTML Templates (`store/templates/store/`)', level=3)
    doc.add_paragraph("These files generate the visual User Interface (UI) using Tailwind CSS.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('base.html: ').bold = True
    p.add_run("The master layout file containing the Navigation bar, Cart badge, and global error/success messages. All other pages 'extend' this file so they don't have to rewrite the navbar.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('home.html: ').bold = True
    p.add_run("The product grid and category sidebar.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('cart.html, orders.html, login.html, signup.html: ').bold = True
    p.add_run("Specific UI pages for their respective functionalities.")
    
    # Template Tags
    doc.add_heading('Custom Template Tags (`store/templatetags/`)', level=3)
    p = doc.add_paragraph()
    p.add_run('cart.py: ').bold = True
    p.add_run("Contains custom Python functions that can be called directly inside HTML templates. For example, calculating the total price of the cart dynamically without needing JavaScript.")

    # Save the document
    doc.save('Project_Documentation.docx')

if __name__ == '__main__':
    create_doc()
